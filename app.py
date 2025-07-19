import os
from flask import Flask, render_template, request, send_file, redirect, url_for,Blueprint,jsonify, flash
from werkzeug.exceptions import HTTPException
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter, ImageDraw, ImageFont, ImageColor
import io
import os
import uuid
import tempfile
from rembg import remove
import pdfplumber
import arabic_reshaper
from bidi.algorithm import get_display
from fpdf import FPDF
from pdf2docx import Converter
from docx import Document
import openpyxl
from werkzeug.utils import secure_filename
import time
from flask_wtf import FlaskForm
from wtforms import StringField, TextAreaField, SubmitField
from wtforms.validators import DataRequired, Email

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'  # Add this line
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///cv_database.db'  # Add this line
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False  # Add this line
app.config['UPLOAD_FOLDER'] = os.path.join('static', 'uploads')
app.config['OUTPUT_FOLDER'] = os.path.join('static', 'outputs')
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'txt', 'docx', 'xlsx', 'png', 'jpg', 'jpeg'}

db.init_app(app)  # Add this line to initialize SQLAlchemy with your Flask app
app.config['UPLOAD_FOLDER'] = os.path.join('static', 'uploads')
app.config['OUTPUT_FOLDER'] = os.path.join('static', 'outputs')
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'txt', 'docx', 'xlsx', 'png', 'jpg', 'jpeg'}

# Ensure upload and output directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

@app.route('/img_editor')
def img_editor():
    return render_template('image-editor.html')

@app.route('/process_image', methods=['POST'])
def process_image():
    if 'filename' not in request.form:
        return jsonify({'error': 'No filename provided'}), 400
        
    filename = request.form.get('filename')
    if not filename:
        return jsonify({'error': 'Empty filename provided'}), 400
        
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    try:
        image = Image.open(input_path).convert('RGBA')
    except FileNotFoundError:
        return jsonify({'error': 'Image not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
    # Get processing parameters
    brightness = int(request.form.get('brightness', 0))
    contrast = int(request.form.get('contrast', 0))
    sharpen = float(request.form.get('sharpen', 0))
    rotate = int(request.form.get('rotate', 0))
    zoom = float(request.form.get('zoom', 1.0))
    
    # Apply transformations
    if rotate != 0:
        image = image.rotate(rotate, expand=True)

    if zoom != 1.0:
        new_size = (int(image.width * zoom), int(image.height * zoom))
        image = image.resize(new_size, Image.LANCZOS)

    if brightness != 0:
        image = apply_brightness(image, brightness)
    
    if contrast != 0:
        image = apply_contrast(image, contrast)
    
    if sharpen > 0:
        image = apply_sharpen(image, sharpen)
    
    # Process watermark
    watermark_text = request.form.get('watermark_text', '')
    if watermark_text:
        watermark_position = request.form.get('watermark_position', 'bottom-right')
        watermark_size = int(request.form.get('watermark_size', 24))
        watermark_color = request.form.get('watermark_color', '#ffffff')
        watermark_rotation = int(request.form.get('watermark_rotation', 0))
        watermark_opacity = float(request.form.get('watermark_opacity', 0.5))
        
        image = add_watermark(
            image=image,
            text=watermark_text,
            position=watermark_position,
            font_size=watermark_size,
            color=watermark_color,
            rotation=watermark_rotation,
            opacity=watermark_opacity
        )
    
    # Handle background removal
    bg_remove = request.form.get('remove_bg', 'false').lower() == 'true'
    if bg_remove:
        image = remove(image)
    
    # Save the processed image
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"processed_{filename}")
    image.save(output_path, format='PNG')
    
    return jsonify({'filename': f"processed_{filename}"})  # Use jsonify here

@app.route('/download_image/<filename>')
def download_image(filename):
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    if not os.path.exists(path):
        return 'File not found', 404
        
    return send_file(path, as_attachment=True)

@app.route('/upload_image', methods=['POST'])
def upload_image():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
        
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
        
    if file:
        try:
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4()}_{file.filename}")
            file.save(input_path)
            
            # Open image for processing
            original = Image.open(input_path).convert('RGBA')
            
            # Save the processed image
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"processed_{os.path.basename(input_path)}")
            original.save(output_path, format='PNG')
            
            return jsonify({'filename': os.path.basename(output_path)})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    return jsonify({'error': 'Upload failed'}), 500

@app.route('/delete_image/<filename>')
def delete_image(filename):
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    try:
        if os.path.exists(path):
            os.remove(path)
            return jsonify({'status': 'deleted'})
        return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# Conversion Functions (Same as before, but with file handling)
def pdf_to_txt(pdf_path, txt_path, use_ocr=False):
    with open(txt_path, 'w', encoding='utf-8') as txt_file:
        if use_ocr:
            try:
                from pdf2image import convert_from_path
                images = convert_from_path(pdf_path)
                for img in images:
                    text = pytesseract.image_to_string(img)
                    txt_file.write(text + "\n")
            except ImportError:
                return False, "pdf2image not installed. Install with: pip install pdf2image"
        else:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    txt_file.write(page.extract_text() + "\n")
    return True, "Success"

def txt_to_pdf(txt_path, pdf_path):
    # Initialize FPDF with UTF-8 encoding to support a wider range of characters
    pdf = FPDF(unit="mm", format="A4")
    # Add a UTF-8 compatible font
    # You might need to download DejaVuSansCondensed.ttf and place it in a 'fonts' directory
    # or provide the full path to the .ttf file.
    pdf.add_font('ArabicTypesetting', '', 'C:/Windows/Fonts/arabtype.ttf', uni=True)
    pdf.add_page()
    pdf.set_font('ArabicTypesetting', '', 12)
    with open(txt_path, 'r', encoding='utf-8') as txt_file:
        for line in txt_file:
            # FPDF's write method should now handle UTF-8 strings correctly
            reshaped_text = arabic_reshaper.reshape(line)
            bidi_text = get_display(reshaped_text)
            pdf.write(8, bidi_text)
    pdf.output(pdf_path)

def pdf_to_docx(pdf_path, docx_path, use_ocr=False):
    if use_ocr:
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(pdf_path)
            doc = Document()
            for img in images:
                text = pytesseract.image_to_string(img)
                doc.add_paragraph(text)
            doc.save(docx_path)
        except ImportError:
            return False, "pdf2image not installed. Install with: pip install pdf2image"
    else:
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
    return True, "Success"

def docx_to_txt(docx_path, txt_path):
    doc = Document(docx_path)
    with open(txt_path, 'w', encoding='utf-8') as txt_file:
        for para in doc.paragraphs:
            txt_file.write(para.text + "\n")

def xls_to_pdf(xls_path, pdf_path):
    # Check if the file has a supported Excel extension
    supported_extensions = ('.xlsx', '.xlsm', '.xltx', '.xltm')
    if not xls_path.lower().endswith(supported_extensions):
        return False, f"Unsupported file format. Please use one of: {', '.join(supported_extensions)}"

    try:
        wb = openpyxl.load_workbook(xls_path)
        sheet = wb.active
        # Initialize FPDF
        pdf = FPDF(unit="mm", format="A4")
        pdf.add_page()
        # Use a font that supports UTF-8 characters, especially for Arabic.
        # Ensure 'arabtype.ttf' is available at the specified path.
        pdf.add_font('ArabicTypesetting', '', 'C:/Windows/Fonts/arabtype.ttf', uni=True)
        pdf.set_font('ArabicTypesetting', '', 12)
        for row in sheet.iter_rows(values_only=True):
            line = " | ".join(str(cell) for cell in row)
            reshaped_text = arabic_reshaper.reshape(line)
            bidi_text = get_display(reshaped_text)
            pdf.cell(200, 10, txt=bidi_text, ln=True)
        pdf.output(pdf_path)
        return True, "Success"
    except Exception as e:
        return False, str(e)


def image_to_txt(image_path, txt_path):
    img = Image.open(image_path)
    text = pytesseract.image_to_string(img,lang="ara+eng")
    with open(txt_path, 'w', encoding='utf-8') as txt_file:
        txt_file.write(text)

app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB


def apply_brightness(image, factor):
    enhancer = ImageEnhance.Brightness(image)
    return enhancer.enhance(1 + factor/100)

def apply_contrast(image, factor):
    enhancer = ImageEnhance.Contrast(image)
    return enhancer.enhance(1 + factor/100)

def apply_sharpen(image, factor):
    if factor == 0:
        return image
    return image.filter(ImageFilter.UnsharpMask(factor))

def add_watermark(image, text, position, font_size, color, rotation, opacity):
    if not text:
        return image
        
    # Create a copy of the image and ensure it's RGBA
    watermarked = image.copy().convert('RGBA')
    
    # Create drawing context
    draw = ImageDraw.Draw(watermarked)
    
    try:
        # Try to use Arial font, fallback to default if not available
        font = ImageFont.truetype("arial.ttf", font_size)
    except:
        font = ImageFont.load_default()
    
    # Get text size using textbbox
    left, top, right, bottom = draw.textbbox((0, 0), text, font=font)
    text_width = right - left
    text_height = bottom - top
    
    # Calculate position
    if position == 'top-left':
        pos = (10, 10)
    elif position == 'top-center':
        pos = (image.width / 2 - text_width / 2, 10)
    elif position == 'top-right':
        pos = (image.width - text_width - 10, 10)
    elif position == 'bottom-left':
        pos = (10, image.height - text_height - 10)
    elif position == 'bottom-center':
        pos = (image.width / 2 - text_width / 2, image.height - text_height - 10)
    elif position == 'bottom-right':
        pos = (image.width - text_width - 10, image.height - text_height - 10)
    elif position == 'center':
        pos = ((image.width - text_width) / 2, (image.height - text_height) / 2)
    else:  # Default to bottom-right if position is not recognized
        pos = (image.width - text_width - 10, image.height - text_height - 10)
    
    # Create a transparent layer for the watermark with same size as original
    watermark_layer = Image.new('RGBA', watermarked.size, (0,0,0,0))
    watermark_draw = ImageDraw.Draw(watermark_layer)
    
    # Convert color from hex to RGB
    rgb_color = ImageColor.getrgb(color)
    
    # Draw the text with rotation
    watermark_draw.text(pos, text, font=font, fill=(*rgb_color, int(255 * opacity)))
    
    if rotation != 0:
        # Rotate the watermark layer while keeping the same size
        watermark_layer = watermark_layer.rotate(rotation, expand=True, center=(pos[0] + text_width / 2, pos[1] + text_height / 2))
        # Recalculate position after rotation
        left, top, right, bottom = watermark_layer.getbbox()
        pos = (pos[0] - left, pos[1] - top)

    # Ensure both images are in RGBA mode before compositing
    watermarked = watermarked.convert('RGBA')
    watermark_layer = watermark_layer.convert('RGBA')
    
    # Composite the watermark onto the original image
    return Image.alpha_composite(watermarked, watermark_layer)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'file' not in request.files:
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            upload_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(upload_path)
            
            conversion_type = request.form.get('conversion_type')
            use_ocr = request.form.get('use_ocr') == 'on'
            
            output_filename_base = f"converted_{filename.rsplit('.', 1)[0]}"
            ext_map = {
                'pdf_to_txt': '.txt',
                'txt_to_pdf': '.pdf',
                'pdf_to_docx': '.docx',
                'docx_to_txt': '.txt',
                'xls_to_pdf': '.pdf',
                'image_to_txt': '.txt'
            }
            ext = ext_map.get(conversion_type, '')
            output_filename = output_filename_base + ext
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            
            start_time = time.time()
            
            success = False
            message = ""
            if conversion_type == 'pdf_to_txt':
                success, message = pdf_to_txt(upload_path, output_path, use_ocr)
            elif conversion_type == 'txt_to_pdf':
                txt_to_pdf(upload_path, output_path)
                success, message = True, "Success"
            elif conversion_type == 'pdf_to_docx':
                success, message = pdf_to_docx(upload_path, output_path, use_ocr)
            elif conversion_type == 'docx_to_txt':
                docx_to_txt(upload_path, output_path)
                success, message = True, "Success"
            elif conversion_type == 'xls_to_pdf':
                success, message = xls_to_pdf(upload_path, output_path)
            elif conversion_type == 'image_to_txt':
                image_to_txt(upload_path, output_path)
                success, message = True, "Success"
            else:
                return "Invalid conversion type", 400
            
            elapsed_time = time.time() - start_time
            
            if success:
                if not os.path.exists(output_path):
                    return render_template('index.html', error=True, message="Conversion failed: output file was not created.")
                return render_template('index.html', 
                                       success=True,
                                       message=f"Conversion successful! Time taken: {elapsed_time:.2f} seconds",
                                       download_file=output_filename)
            else:
                return render_template('index.html', 
                                       error=True,
                                       message=message)
    
    return render_template('index.html')
@app.route('/canvas')
def canvas():
	return render_template('canvas.html')

@app.route('/pdfimg')
def pdfimg():
	return render_template('pdf-to-img-img-to-pdf.html')	


@app.route('/download/<filename>')
def download(filename):
    # Secure the filename to prevent path traversal
    filename = secure_filename(filename)
    file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    if not os.path.exists(file_path):
        return render_template('index.html', error=True, message="File not found for download.")
    return send_file(file_path, as_attachment=True)



import traceback

@app.errorhandler(Exception)
def handle_exception(e):
    # pass through HTTP errors
    if isinstance(e, HTTPException):
        return e

    # handle non-HTTP errors
    error_type = type(e).__name__
    error_message = str(e)
    
    # Get traceback only if debug mode is on
    tb = traceback.format_exc() if app.debug else None

    return render_template("error.html", 
                           error_type=error_type, 
                           error_message=error_message, 
                           traceback=tb), 500

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)
